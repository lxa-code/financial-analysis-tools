import os, json

# 读取日报
print("=" * 60)
print("【日报文件】")
print("=" * 60)
docx_path = r"e:\BaiduSyncdisk\CODE\城发投资工作专区\10_日常工作\个人工作日报\新建 DOCX 文档.docx"
try:
    import docx
    doc = docx.Document(docx_path)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  [{i}] {para.text.strip()}")
    # 也读取表格
    for ti, table in enumerate(doc.tables):
        print(f"\n  --- 表格 {ti+1} ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"    Row{ri}: {cells}")
except Exception as e:
    print(f"  错误: {e}")

# 读取周报
print("\n" + "=" * 60)
print("【周报文件 - 个人周报_20260525-0527.docx】")
print("=" * 60)
docx_path2 = r"e:\BaiduSyncdisk\CODE\城发投资工作专区\10_日常工作\个人工作周报\个人周报_20260525-0527.docx"
try:
    doc2 = docx.Document(docx_path2)
    for i, para in enumerate(doc2.paragraphs):
        if para.text.strip():
            print(f"  [{i}] {para.text.strip()}")
    for ti, table in enumerate(doc2.tables):
        print(f"\n  --- 表格 {ti+1} ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"    Row{ri}: {cells}")
except Exception as e:
    print(f"  错误: {e}")

# 读取周报模板
print("\n" + "=" * 60)
print("【周报模板】")
print("=" * 60)
docx_path3 = r"e:\BaiduSyncdisk\CODE\城发投资工作专区\10_日常工作\个人工作周报\个人周报模板.docx"
try:
    doc3 = docx.Document(docx_path3)
    for i, para in enumerate(doc3.paragraphs):
        if para.text.strip():
            print(f"  [{i}] {para.text.strip()}")
    for ti, table in enumerate(doc3.tables):
        print(f"\n  --- 表格 {ti+1} ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"    Row{ri}: {cells}")
except Exception as e:
    print(f"  错误: {e}")

# 读取周例会材料模板
print("\n" + "=" * 60)
print("【周例会材料要素组模板】")
print("=" * 60)
xlsx_path = r"e:\BaiduSyncdisk\CODE\城发投资工作专区\10_日常工作\个人工作周报\2-1 城发投资周例会材料要素组模板.xlsx"
try:
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path)
    for sn in wb.sheetnames:
        ws = wb[sn]
        print(f"\n  Sheet: {sn} (rows={ws.max_row}, cols={ws.max_column})")
        for ri in range(1, min(ws.max_row + 1, 50)):
            row_data = []
            for ci in range(1, ws.max_column + 1):
                v = ws.cell(row=ri, column=ci).value
                if v is not None:
                    row_data.append(f"C{ci}:{v}")
            if row_data:
                print(f"    Row{ri}: {row_data}")
except Exception as e:
    print(f"  错误: {e}")

# 读取模板文件
print("\n" + "=" * 60)
print("【个人周工作总结计划模板.xlsx】")
print("=" * 60)
template_path = r"D:\桌面\桌面\个人周工作总结计划模板.xlsx"
try:
    wb2 = openpyxl.load_workbook(template_path)
    for sn in wb2.sheetnames:
        ws2 = wb2[sn]
        print(f"\n  Sheet: {sn} (rows={ws2.max_row}, cols={ws2.max_column})")
        for ri in range(1, ws2.max_row + 1):
            row_data = []
            for ci in range(1, ws2.max_column + 1):
                v = ws2.cell(row=ri, column=ci).value
                if v is not None:
                    row_data.append(f"C{ci}:{v}")
            if row_data:
                print(f"    Row{ri}: {row_data}")
except Exception as e:
    print(f"  错误: {e}")
