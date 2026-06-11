#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成WORD文字说明模板（带数据占位符）
基于：城发投资月度财务情况说明模板
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

def set_run_font(run, size=16, bold=False, name='仿宋_GB2312', color_rgb=None):
    """设置run的字体格式"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)

def add_para(doc, text, font_size=16, bold=False, font_name='仿宋_GB2312', 
              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=30, 
              space_before=0, space_after=0):
    """添加段落并设置格式"""
    para = doc.add_paragraph()
    para.alignment = alignment
    para_format = para.paragraph_format
    para_format.line_spacing = Pt(line_spacing)
    para_format.space_before = Pt(space_before)
    para_format.space_after = Pt(space_after)
    
    run = para.add_run(text)
    set_run_font(run, size=font_size, bold=bold, name=font_name)
    return para

def add_table(doc, headers, data_rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=12, bold=True, name='黑体')
    
    # 数据行
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_data in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(cell_data)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=11, name='仿宋_GB2312')
    
    doc.add_paragraph()  # 表格后空行

def generate():
    doc = Document()
    
    # ========== 封面 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('重要资料\n注意保密')
    set_run_font(run, size=16, bold=True, name='黑体', color_rgb=(0,0,0))
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('城发投资{{YYYY}}年度{{MM}}月份\n月度财务情况说明')
    set_run_font(run, size=22, bold=False, name='方正小标宋简体', color_rgb=(0,0,0))
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('二零二六年{{MM}}月')
    set_run_font(run, size=16, bold=False, name='仿宋_GB2312', color_rgb=(0,0,0))
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    add_para(doc, '目录', font_size=22, bold=False, font_name='方正小标宋简体', 
             line_spacing=30, space_before=0, space_after=0)
    
    toc_items = [
        '一、关键指标完成情况\t2',
        '（一）整体情况\t3',
        '（利润总额、管理费用、回款情况）',
        '（二）本部情况\t3',
        '（利润总额、管理费用、资金情况）',
        '二、盈利情况\t3',
        '（一）整体盈利情况\t3',
        '（二）本部盈利情况\t4',
        '（三）下属企业盈利情况\t4',
        '三、资产负债情况\t5',
        '（一）整体情况\t5',
        '（二）本部情况\t6',
        '四、资金情况\t9',
        '（一）整体情况\t9',
        '（二）本部情况\t9',
        '（三）归集情况\t10',
        '五、投资情况\t11',
        '（一）股权投资情况\t11',
        '（二）债权投资情况\t11',
        '六、预算情况\t11',
        '（一）预算执行情况\t11',
        '（利润总额、管理费用、回款情况）',
        '（二）预算计划情况\t11'
    ]
    for item in toc_items:
        add_para(doc, item, font_size=16, bold=False, font_name='仿宋_GB2312',
                 line_spacing=30, space_before=0, space_after=0)
    
    doc.add_page_break()
    
    # ========== 一、关键指标完成情况 ==========
    add_para(doc, '一、关键指标完成情况', font_size=22, bold=False, 
             font_name='方正小标宋简体', line_spacing=30)
    
    # （一）整体情况
    add_para(doc, '（一）整体情况', font_size=16, bold=False, 
             font_name='黑体', line_spacing=30)
    add_para(doc, '（利润总额、管理费用、回款情况）', font_size=16, bold=False,
             font_name='仿宋_GB2312', line_spacing=30)
    
    txt1 = ('截至{{YYYY}}年{{MM}}月，城发投资整体实现利润总额{{利润总额_整体}}万元，'
             '同比增长{{利润总额_整体_同比}}%；管理费用{{管理费用_整体}}万元，'
             '同比下降{{管理费用_整体_同比}}%；回款金额{{回款金额_整体}}万元，'
             '同比增长{{回款金额_整体_同比}}%。')
    add_para(doc, txt1, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    add_table(doc, 
        headers=['指标', '{{YYYY}}年{{MM}}月', '上年同期', '同比增减', '同比增减率'],
        data_rows=[
            ['利润总额（万元）', '{{利润总额_整体}}', '{{利润总额_整体_上年}}', 
             '{{利润总额_整体_同比增减}}', '{{利润总额_整体_同比}}%'],
            ['管理费用（万元）', '{{管理费用_整体}}', '{{管理费用_整体_上年}}', 
             '{{管理费用_整体_同比增减}}', '{{管理费用_整体_同比}}%'],
            ['回款金额（万元）', '{{回款金额_整体}}', '{{回款金额_整体_上年}}', 
             '{{回款金额_整体_同比增减}}', '{{回款金额_整体_同比}}%']
        ]
    )
    
    # （二）本部情况
    add_para(doc, '（二）本部情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    add_para(doc, '（利润总额、管理费用、资金情况）', font_size=16, bold=False,
             font_name='仿宋_GB2312', line_spacing=30)
    
    txt2 = ('截至{{YYYY}}年{{MM}}月，城发投资本部实现利润总额{{利润总额_本部}}万元，'
             '同比增长{{利润总额_本部_同比}}%；管理费用{{管理费用_本部}}万元，'
             '同比下降{{管理费用_本部_同比}}%；资金余额{{资金余额_本部}}万元。')
    add_para(doc, txt2, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    add_table(doc,
        headers=['指标', '{{YYYY}}年{{MM}}月', '上年同期', '同比增减', '同比增减率'],
        data_rows=[
            ['利润总额（万元）', '{{利润总额_本部}}', '{{利润总额_本部_上年}}',
             '{{利润总额_本部_同比增减}}', '{{利润总额_本部_同比}}%'],
            ['管理费用（万元）', '{{管理费用_本部}}', '{{管理费用_本部_上年}}',
             '{{管理费用_本部_同比增减}}', '{{管理费用_本部_同比}}%'],
            ['资金余额（万元）', '{{资金余额_本部}}', '{{资金余额_本部_上年}}',
             '{{资金余额_本部_同比增减}}', '-']
        ]
    )
    
    doc.add_page_break()
    
    # ========== 二、盈利情况 ==========
    add_para(doc, '二、盈利情况', font_size=22, bold=False,
             font_name='方正小标宋简体', line_spacing=30)
    
    # （一）整体盈利情况
    add_para(doc, '（一）整体盈利情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt3 = ('截至{{YYYY}}年{{MM}}月，城发投资整体实现营业收入{{营业收入_整体}}万元，'
             '同比增长{{营业收入_整体_同比}}%；利润总额{{利润总额_整体}}万元，'
             '同比增长{{利润总额_整体_同比}}%；净利润{{净利润_整体}}万元，'
             '同比增长{{净利润_整体_同比}}%。')
    add_para(doc, txt3, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    add_table(doc,
        headers=['项目', '{{YYYY}}年{{MM}}月', '上年同期', '同比增减', '同比增减率'],
        data_rows=[
            ['营业收入（万元）', '{{营业收入_整体}}', '{{营业收入_整体_上年}}',
             '{{营业收入_整体_同比增减}}', '{{营业收入_整体_同比}}%'],
            ['营业成本（万元）', '{{营业成本_整体}}', '{{营业成本_整体_上年}}',
             '{{营业成本_整体_同比增减}}', '{{营业成本_整体_同比}}%'],
            ['利润总额（万元）', '{{利润总额_整体}}', '{{利润总额_整体_上年}}',
             '{{利润总额_整体_同比增减}}', '{{利润总额_整体_同比}}%'],
            ['净利润（万元）', '{{净利润_整体}}', '{{净利润_整体_上年}}',
             '{{净利润_整体_同比增减}}', '{{净利润_整体_同比}}%']
        ]
    )
    
    # （二）本部盈利情况
    add_para(doc, '（二）本部盈利情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt4 = ('截至{{YYYY}}年{{MM}}月，城发投资本部实现营业收入{{营业收入_本部}}万元，'
             '同比增长{{营业收入_本部_同比}}%；利润总额{{利润总额_本部}}万元，'
             '同比增长{{利润总额_本部_同比}}%。')
    add_para(doc, txt4, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    # （三）下属企业盈利情况
    add_para(doc, '（三）下属企业盈利情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt5 = '截至{{YYYY}}年{{MM}}月，主要下属企业盈利情况如下：'
    add_para(doc, txt5, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    add_table(doc,
        headers=['企业名称', '营业收入（万元）', '利润总额（万元）', '净利润（万元）', '同比增减率'],
        data_rows=[
            ['{{企业1名称}}', '{{企业1_营业收入}}', '{{企业1_利润总额}}', 
             '{{企业1_净利润}}', '{{企业1_利润同比}}%'],
            ['{{企业2名称}}', '{{企业2_营业收入}}', '{{企业2_利润总额}}', 
             '{{企业2_净利润}}', '{{企业2_利润同比}}%'],
            ['（其他企业）', '...', '...', '...', '...']
        ]
    )
    
    doc.add_page_break()
    
    # ========== 三、资产负债情况 ==========
    add_para(doc, '三、资产负债情况', font_size=22, bold=False,
             font_name='方正小标宋简体', line_spacing=30)
    
    # （一）整体情况
    add_para(doc, '（一）整体情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt6 = ('截至{{YYYY}}年{{MM}}月末，城发投资总资产{{总资产_整体}}万元，'
             '总负债{{总负债_整体}}万元，所有者权益{{所有者权益_整体}}万元，'
             '资产负债率{{资产负债率_整体}}%。')
    add_para(doc, txt6, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    add_table(doc,
        headers=['项目', '{{YYYY}}年{{MM}}月末', '上年末', '增减额', '增减率'],
        data_rows=[
            ['资产总计（万元）', '{{总资产_整体}}', '{{总资产_整体_上年末}}',
             '{{总资产_整体_增减}}', '{{总资产_整体_增减率}}%'],
            ['负债合计（万元）', '{{总负债_整体}}', '{{总负债_整体_上年末}}',
             '{{总负债_整体_增减}}', '{{总负债_整体_增减率}}%'],
            ['所有者权益（万元）', '{{所有者权益_整体}}', '{{所有者权益_整体_上年末}}',
             '{{所有者权益_整体_增减}}', '{{所有者权益_整体_增减率}}%'],
            ['资产负债率（%）', '{{资产负债率_整体}}', '{{资产负债率_整体_上年末}}', '-', '-']
        ]
    )
    
    # （二）本部情况
    add_para(doc, '（二）本部情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt7 = ('截至{{YYYY}}年{{MM}}月末，城发投资本部总资产{{总资产_本部}}万元，'
             '总负债{{总负债_本部}}万元，资产负债率{{资产负债率_本部}}%。')
    add_para(doc, txt7, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    doc.add_page_break()
    
    # ========== 四、资金情况 ==========
    add_para(doc, '四、资金情况', font_size=22, bold=False,
             font_name='方正小标宋简体', line_spacing=30)
    
    # （一）整体情况
    add_para(doc, '（一）整体情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt8 = ('截至{{YYYY}}年{{MM}}月末，城发投资整体资金余额{{资金余额_整体}}万元，'
             '其中货币资金{{货币资金_整体}}万元，其他流动资金{{其他流动资金_整体}}万元。')
    add_para(doc, txt8, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    # （二）本部情况
    add_para(doc, '（二）本部情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt9 = ('截至{{YYYY}}年{{MM}}月末，城发投资本部资金余额{{资金余额_本部}}万元，'
             '较上月{{资金余额_本部_环比增减}}万元。')
    add_para(doc, txt9, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    # （三）归集情况
    add_para(doc, '（三）归集情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt10 = ('截至{{YYYY}}年{{MM}}月末，城发投资资金归集金额{{归集金额}}万元，'
              '归集率{{归集率}}%。')
    add_para(doc, txt10, font_size=16, bold=False, font_name='仿宋_GB2312',
              line_spacing=30)
    
    add_table(doc,
        headers=['企业名称', '资金余额（万元）', '归集金额（万元）', '归集率（%）'],
        data_rows=[
            ['{{企业1名称}}', '{{企业1_资金余额}}', '{{企业1_归集金额}}', '{{企业1_归集率}}'],
            ['{{企业2名称}}', '{{企业2_资金余额}}', '{{企业2_归集金额}}', '{{企业2_归集率}}'],
            ['（其他企业）', '...', '...', '...']
        ]
    )
    
    doc.add_page_break()
    
    # ========== 五、投资情况 ==========
    add_para(doc, '五、投资情况', font_size=22, bold=False,
             font_name='方正小标宋简体', line_spacing=30)
    
    # （一）股权投资情况
    add_para(doc, '（一）股权投资情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt11 = ('截至{{YYYY}}年{{MM}}月末，城发投资股权投资总额{{股权投资总额}}万元，'
              '其中本年新增{{本年新增股权投资}}万元。')
    add_para(doc, txt11, font_size=16, bold=False, font_name='仿宋_GB2312',
              line_spacing=30)
    
    # （二）债权投资情况
    add_para(doc, '（二）债权投资情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt12 = ('截至{{YYYY}}年{{MM}}月末，城发投资债权投资总额{{债权投资总额}}万元，'
              '其中委托贷款{{委托贷款总额}}万元。')
    add_para(doc, txt12, font_size=16, bold=False, font_name='仿宋_GB2312',
              line_spacing=30)
    
    doc.add_page_break()
    
    # ========== 六、预算情况 ==========
    add_para(doc, '六、预算情况', font_size=22, bold=False,
             font_name='方正小标宋简体', line_spacing=30)
    
    # （一）预算执行情况
    add_para(doc, '（一）预算执行情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    add_para(doc, '（利润总额、管理费用、回款情况）', font_size=16, bold=False,
             font_name='仿宋_GB2312', line_spacing=30)
    
    txt13 = ('截至{{YYYY}}年{{MM}}月，城发投资累计完成利润总额{{利润总额_执行}}万元，'
              '占年度预算的{{利润总额_执行率}}%；管理费用{{管理费用_执行}}万元，'
              '占年度预算的{{管理费用_执行率}}%；回款金额{{回款_执行}}万元，'
              '占年度预算的{{回款_执行率}}%。')
    add_para(doc, txt13, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    add_table(doc,
        headers=['项目', '年度预算（万元）', '实际完成（万元）', '执行率（%）', '序时进度（%）'],
        data_rows=[
            ['利润总额', '{{利润总额_预算}}', '{{利润总额_执行}}', 
             '{{利润总额_执行率}}', '{{利润总额_序时进度}}'],
            ['管理费用', '{{管理费用_预算}}', '{{管理费用_执行}}', 
             '{{管理费用_执行率}}', '{{管理费用_序时进度}}'],
            ['回款金额', '{{回款_预算}}', '{{回款_执行}}', 
             '{{回款_执行率}}', '{{回款_序时进度}}']
        ]
    )
    
    # （二）预算计划情况
    add_para(doc, '（二）预算计划情况', font_size=16, bold=False,
             font_name='黑体', line_spacing=30)
    
    txt14 = ('根据{{YYYY}}年度预算安排，后续月份需完成利润总额{{利润总额_待完成}}万元，'
              '管理费用控制在{{管理费用_待完成}}万元以内。')
    add_para(doc, txt14, font_size=16, bold=False, font_name='仿宋_GB2312',
             line_spacing=30)
    
    # ========== 保存 ==========
    output = r'c:\Users\lecoo\CodeBuddy\20260425100224\城发投资月度财务情况说明_文字模板_带占位符.docx'
    doc.save(output)
    print(f'已生成WORD文字模板：{output}')
    print()
    print('模板中包含的占位符（部分示例）：')
    print('  {{YYYY}}、{{MM}} - 年份和月份')
    print('  {{利润总额_整体}}、{{管理费用_整体}} - 关键指标')
    print('  {{营业收入_整体}}、{{净利润_整体}} - 盈利指标')
    print('  {{总资产_整体}}、{{资产负债率_整体}} - 资产负债指标')
    print('  {{资金余额_整体}}、{{归集率}} - 资金指标')
    print('  {{股权投资总额}}、{{债权投资总额}} - 投资指标')
    print('  {{利润总额_预算}}、{{利润总额_执行率}} - 预算指标')
    print()
    print('下一步：生成EXCEL数据填充模板...')

if __name__ == '__main__':
    generate()
