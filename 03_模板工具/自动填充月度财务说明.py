#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
自动填充月度财务情况说明
读取EXCEL数据 → 填充WORD模板中的占位符 → 生成最终报告
"""

from docx import Document
from openpyxl import load_workbook
import os
import re

def read_excel_data(excel_path):
    """读取EXCEL模板中的所有数据"""
    wb = load_workbook(excel_path, data_only=True)
    data = {}
    
    # Sheet1: 参数设置
    ws = wb['参数设置']
    for row in ws.iter_rows(min_row=3, max_row=6, values_only=True):
        if row[0] and row[1]:
            key = str(row[0]).strip()
            value = row[1]
            if key == 'YYYY':
                data['YYYY'] = str(value)
            elif key == 'MM':
                data['MM'] = str(value).zfill(2)
            elif key == '公司名称':
                data['公司名称'] = str(value)
            elif key == '报告期间':
                data['报告期间'] = str(value)
    
    # Sheet2: 关键指标_整体
    ws = wb['关键指标_整体']
    for row in ws.iter_rows(min_row=3, max_row=5, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '利润总额' in key:
                data['利润总额_整体'] = format_num(row[1])
                data['利润总额_整体_上年'] = format_num(row[2])
                data['利润总额_整体_同比增减'] = format_num(row[3])
                data['利润总额_整体_同比'] = format_num(row[4], is_pct=True)
            elif '管理费用' in key:
                data['管理费用_整体'] = format_num(row[1])
                data['管理费用_整体_上年'] = format_num(row[2])
                data['管理费用_整体_同比增减'] = format_num(row[3])
                data['管理费用_整体_同比'] = format_num(row[4], is_pct=True)
            elif '回款' in key:
                data['回款金额_整体'] = format_num(row[1])
                data['回款金额_整体_上年'] = format_num(row[2])
                data['回款金额_整体_同比增减'] = format_num(row[3])
                data['回款金额_整体_同比'] = format_num(row[4], is_pct=True)
    
    # Sheet3: 关键指标_本部
    ws = wb['关键指标_本部']
    for row in ws.iter_rows(min_row=3, max_row=5, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '利润总额' in key:
                data['利润总额_本部'] = format_num(row[1])
                data['利润总额_本部_上年'] = format_num(row[2])
                data['利润总额_本部_同比增减'] = format_num(row[3])
                data['利润总额_本部_同比'] = format_num(row[4], is_pct=True)
            elif '管理费用' in key:
                data['管理费用_本部'] = format_num(row[1])
                data['管理费用_本部_上年'] = format_num(row[2])
                data['管理费用_本部_同比增减'] = format_num(row[3])
                data['管理费用_本部_同比'] = format_num(row[4], is_pct=True)
            elif '资金余额' in key:
                data['资金余额_本部'] = format_num(row[1])
                data['资金余额_本部_上年'] = format_num(row[2])
                data['资金余额_本部_同比增减'] = format_num(row[3])
    
    # Sheet4: 盈利情况_整体
    ws = wb['盈利情况_整体']
    for row in ws.iter_rows(min_row=3, max_row=6, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '营业收入' in key:
                data['营业收入_整体'] = format_num(row[1])
                data['营业收入_整体_上年'] = format_num(row[2])
                data['营业收入_整体_同比增减'] = format_num(row[3])
                data['营业收入_整体_同比'] = format_num(row[4], is_pct=True)
            elif '营业成本' in key:
                data['营业成本_整体'] = format_num(row[1])
                data['营业成本_整体_上年'] = format_num(row[2])
                data['营业成本_整体_同比增减'] = format_num(row[3])
                data['营业成本_整体_同比'] = format_num(row[4], is_pct=True)
            elif '利润总额' in key and '本部' not in key:
                data['利润总额_整体'] = format_num(row[1])
                data['利润总额_整体_上年'] = format_num(row[2])
                data['利润总额_整体_同比增减'] = format_num(row[3])
                data['利润总额_整体_同比'] = format_num(row[4], is_pct=True)
            elif '净利润' in key:
                data['净利润_整体'] = format_num(row[1])
                data['净利润_整体_上年'] = format_num(row[2])
                data['净利润_整体_同比增减'] = format_num(row[3])
                data['净利润_整体_同比'] = format_num(row[4], is_pct=True)
    
    # Sheet5: 盈利情况_本部
    ws = wb['盈利情况_本部']
    for row in ws.iter_rows(min_row=3, max_row=4, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '营业收入' in key:
                data['营业收入_本部'] = format_num(row[1])
                data['营业收入_本部_上年'] = format_num(row[2])
                data['营业收入_本部_同比增减'] = format_num(row[3])
                data['营业收入_本部_同比'] = format_num(row[4], is_pct=True)
            elif '利润总额' in key:
                data['利润总额_本部'] = format_num(row[1])
                data['利润总额_本部_上年'] = format_num(row[2])
                data['利润总额_本部_同比增减'] = format_num(row[3])
                data['利润总额_本部_同比'] = format_num(row[4], is_pct=True)
    
    # Sheet7: 资产负债_整体
    ws = wb['资产负债_整体']
    for row in ws.iter_rows(min_row=3, max_row=6, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '资产总计' in key:
                data['总资产_整体'] = format_num(row[1])
                data['总资产_整体_上年末'] = format_num(row[2])
                data['总资产_整体_增减'] = format_num(row[3])
                data['总资产_整体_增减率'] = format_num(row[4], is_pct=True)
            elif '负债合计' in key:
                data['总负债_整体'] = format_num(row[1])
                data['总负债_整体_上年末'] = format_num(row[2])
                data['总负债_整体_增减'] = format_num(row[3])
                data['总负债_整体_增减率'] = format_num(row[4], is_pct=True)
            elif '所有者权益' in key:
                data['所有者权益_整体'] = format_num(row[1])
                data['所有者权益_整体_上年末'] = format_num(row[2])
                data['所有者权益_整体_增减'] = format_num(row[3])
                data['所有者权益_整体_增减率'] = format_num(row[4], is_pct=True)
            elif '资产负债率' in key and '本部' not in key:
                data['资产负债率_整体'] = format_num(row[1], is_pct=True)
                data['资产负债率_整体_上年末'] = format_num(row[2], is_pct=True)
    
    # Sheet8: 资产负债_本部
    ws = wb['资产负债_本部']
    for row in ws.iter_rows(min_row=3, max_row=5, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '资产总计' in key:
                data['总资产_本部'] = format_num(row[1])
            elif '负债合计' in key:
                data['总负债_本部'] = format_num(row[1])
            elif '资产负债率' in key:
                data['资产负债率_本部'] = format_num(row[1], is_pct=True)
    
    # Sheet9: 资金情况_整体
    ws = wb['资金情况_整体']
    for row in ws.iter_rows(min_row=3, max_row=5, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '资金余额' in key and '本部' not in key:
                data['资金余额_整体'] = format_num(row[1])
            elif '货币资金' in key:
                data['货币资金_整体'] = format_num(row[1])
            elif '其他流动资金' in key:
                data['其他流动资金_整体'] = format_num(row[1])
    
    # Sheet10: 资金情况_本部
    ws = wb['资金情况_本部']
    for row in ws.iter_rows(min_row=3, max_row=3, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '资金余额' in key:
                data['资金余额_本部'] = format_num(row[1])
                data['资金余额_本部_上月'] = format_num(row[2])
                data['资金余额_本部_环比增减'] = format_num(row[3])
    
    # Sheet12: 投资情况
    ws = wb['投资情况']
    for row in ws.iter_rows(min_row=3, max_row=6, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '股权投资总额' in key and '本年新增' not in key:
                data['股权投资总额'] = format_num(row[1])
            elif '本年新增股权投资' in key:
                data['本年新增股权投资'] = format_num(row[1])
            elif '债权投资总额' in key:
                data['债权投资总额'] = format_num(row[1])
            elif '委托贷款总额' in key:
                data['委托贷款总额'] = format_num(row[1])
    
    # Sheet13: 预算执行
    ws = wb['预算执行']
    for row in ws.iter_rows(min_row=3, max_row=5, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '利润总额' in key and '待完成' not in key:
                data['利润总额_预算'] = format_num(row[1])
                data['利润总额_执行'] = format_num(row[2])
                data['利润总额_执行率'] = format_num(row[3], is_pct=True)
                data['利润总额_序时进度'] = format_num(row[4], is_pct=True)
            elif '管理费用' in key and '待完成' not in key:
                data['管理费用_预算'] = format_num(row[1])
                data['管理费用_执行'] = format_num(row[2])
                data['管理费用_执行率'] = format_num(row[3], is_pct=True)
                data['管理费用_序时进度'] = format_num(row[4], is_pct=True)
            elif '回款' in key:
                data['回款_预算'] = format_num(row[1])
                data['回款_执行'] = format_num(row[2])
                data['回款_执行率'] = format_num(row[3], is_pct=True)
                data['回款_序时进度'] = format_num(row[4], is_pct=True)
    
    # Sheet14: 预算计划
    ws = wb['预算计划']
    for row in ws.iter_rows(min_row=3, max_row=4, values_only=True):
        if row[0]:
            key = str(row[0]).strip()
            if '利润总额待完成' in key:
                data['利润总额_待完成'] = format_num(row[1])
            elif '管理费用待完成' in key:
                data['管理费用_待完成'] = format_num(row[1])
    
    wb.close()
    print(f'已读取EXCEL数据，共{len(data)}个字段')
    return data

def format_num(value, is_pct=False):
    """格式化数字：保留两位小数，添加千分位"""
    if value is None:
        return ''
    try:
        num = float(value)
        if is_pct:
            return f"{num:.2f}"
        else:
            return f"{num:,.2f}"
    except (ValueError, TypeError):
        return str(value) if value else ''

def replace_placeholders_in_paragraph(para, data):
    """替换段落中的占位符"""
    for run in para.runs:
        text = run.text
        new_text = text
        for key, value in data.items():
            placeholder = '{{' + key + '}}'
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))
        
        if new_text != text:
            run.text = new_text

def replace_placeholders_in_cell(cell, data):
    """替换表格单元格中的占位符"""
    for para in cell.paragraphs:
        replace_placeholders_in_paragraph(para, data)

def fill_word_template(word_template_path, excel_data, output_path):
    """填充WORD模板"""
    doc = Document(word_template_path)
    
    # 替换段落中的占位符
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, excel_data)
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders_in_cell(cell, excel_data)
    
    doc.save(output_path)
    print(f'已生成最终报告：{output_path}')

def main():
    """主函数"""
    # 文件路径
    excel_template = r'c:\Users\lecoo\CodeBuddy\20260425100224\城发投资月度财务数据填充模板.xlsx'
    word_template = r'c:\Users\lecoo\CodeBuddy\20260425100224\城发投资月度财务情况说明_文字模板_带占位符.docx'
    output = r'c:\Users\lecoo\CodeBuddy\20260425100224\城发投资月度财务情况说明_最终版.docx'
    
    # 检查文件是否存在
    if not os.path.exists(excel_template):
        print(f'错误：EXCEL模板不存在 - {excel_template}')
        return
    if not os.path.exists(word_template):
        print(f'错误：WORD模板不存在 - {word_template}')
        return
    
    # 读取EXCEL数据
    print('正在读取EXCEL数据...')
    data = read_excel_data(excel_template)
    
    # 填充WORD模板
    print('正在填充WORD模板...')
    fill_word_template(word_template, data, output)
    
    print()
    print('='*50)
    print('完成！生成的文件：')
    print(f'  {output}')
    print('='*50)

if __name__ == '__main__':
    main()
