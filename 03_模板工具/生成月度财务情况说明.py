#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成城发投资月度财务情况说明
根据Word模板结构和Excel数据表生成报告
"""

import os
import sys
from docx import Document
import pandas as pd
import openpyxl
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

def read_word_template(template_path):
    """读取Word模板结构"""
    print(f"正在读取模板文档: {template_path}")
    doc = Document(template_path)
    
    # 提取所有段落文本
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({
                'text': para.text,
                'style': para.style.name if para.style else 'Normal',
                'level': para.style.name.count('Heading') if 'Heading' in para.style.name else 0
            })
    
    # 提取所有表格
    tables = []
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append({
            'index': i,
            'data': table_data,
            'rows': len(table.rows),
            'cols': len(table.columns)
        })
    
    print(f"模板包含 {len(paragraphs)} 个段落, {len(tables)} 个表格")
    return {'paragraphs': paragraphs, 'tables': tables, 'document': doc}

def read_excel_data(excel_path):
    """读取Excel数据表的所有sheet"""
    print(f"正在读取Excel数据: {excel_path}")
    
    # 使用openpyxl读取所有sheet名称
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    sheet_names = wb.sheetnames
    print(f"Excel包含 {len(sheet_names)} 个sheet: {', '.join(sheet_names)}")
    
    # 读取每个sheet的数据
    all_data = {}
    for sheet_name in sheet_names:
        df = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)
        all_data[sheet_name] = df
        print(f"  Sheet '{sheet_name}': {df.shape[0]}行 x {df.shape[1]}列")
    
    return all_data, sheet_names

def analyze_template_structure(template_info):
    """分析模板结构，提取章节和表格模板"""
    print("\n" + "="*60)
    print("模板结构分析")
    print("="*60)
    
    paragraphs = template_info['paragraphs']
    
    # 提取标题层级结构
    print("\n文档结构:")
    for i, para in enumerate(paragraphs[:50]):  # 只显示前50个段落
        if para['text'].strip():
            indent = "  " * (para['level'] - 1) if para['level'] > 0 else ""
            print(f"{indent}{para['text'][:80]}")
    
    # 分析表格结构
    print(f"\n表格数量: {len(template_info['tables'])}")
    for table_info in template_info['tables'][:5]:  # 只显示前5个表格
        print(f"\n表格 {table_info['index'] + 1}: {table_info['rows']}行 x {table_info['cols']}列")
        for row_idx, row_data in enumerate(table_info['data'][:5]):  # 只显示前5行
            print(f"  行{row_idx + 1}: {' | '.join([str(cell)[:15] for cell in row_data])}")
    
    return template_info

def analyze_excel_data(all_data, sheet_names):
    """分析Excel数据内容"""
    print("\n" + "="*60)
    print("Excel数据分析")
    print("="*60)
    
    for sheet_name in sheet_names:
        df = all_data[sheet_name]
        print(f"\nSheet: {sheet_name}")
        print(f"  形状: {df.shape[0]}行 x {df.shape[1]}列")
        
        # 显示前几行数据
        print("  前5行数据:")
        for i in range(min(5, df.shape[0])):
            row_data = df.iloc[i, :].tolist()
            print(f"    行{i+1}: {[str(cell)[:20] for cell in row_data[:10]]}")  # 只显示前10列
        
        # 查找可能的标题行
        for i in range(min(10, df.shape[0])):
            row_data = df.iloc[i, :].tolist()
            # 如果第一行包含"公司"或"指标"等关键词，可能是标题行
            row_str = ' '.join([str(cell) for cell in row_data if pd.notna(cell)])
            if any(keyword in row_str for keyword in ['公司', '指标', '项目', '名称', '日期']):
                print(f"  可能的标题行: 第{i+1}行")
                break

def create_monthly_report(template_path, excel_path, output_path):
    """生成月度财务情况说明"""
    print("\n" + "="*60)
    print("开始生成月度财务情况说明")
    print("="*60)
    
    # 1. 读取模板
    template_info = read_word_template(template_path)
    
    # 2. 读取Excel数据
    all_data, sheet_names = read_excel_data(excel_path)
    
    # 3. 分析模板和数据结构
    analyze_template_structure(template_info)
    analyze_excel_data(all_data, sheet_names)
    
    # 4. 创建新文档（基于模板）
    print("\n正在创建新文档...")
    new_doc = Document(template_path)
    
    # 5. 保存文档
    new_doc.save(output_path)
    print(f"\n文档已保存到: {output_path}")
    print("请手动打开文档，根据实际数据填充内容")
    
    return {
        'template_info': template_info,
        'excel_data': all_data,
        'output_path': output_path
    }

if __name__ == "__main__":
    # 文件路径
    template_path = r"E:\BaiduSyncdisk\城发投\2基础资料\2026年财务分析\每月报送\城发投资月度财务情况说明模板.docx"
    excel_path = r"E:\BaiduSyncdisk\城发投\2基础资料\2026年财务分析\每月报送\每月财务信息输出表 5月.xlsx"
    output_path = r"c:\Users\lecoo\CodeBuddy\20260425100224\城发投资月度财务情况说明_202605.docx"
    
    # 检查文件是否存在
    if not os.path.exists(template_path):
        print(f"错误: 模板文件不存在: {template_path}")
        sys.exit(1)
    
    if not os.path.exists(excel_path):
        print(f"错误: Excel文件不存在: {excel_path}")
        sys.exit(1)
    
    # 生成报告
    result = create_monthly_report(template_path, excel_path, output_path)
    
    print("\n" + "="*60)
    print("处理完成!")
    print("="*60)
    print(f"模板结构已分析，Excel数据已读取")
    print(f"新文档已创建: {output_path}")
    print("\n下一步:")
    print("1. 查看模板结构和Excel数据")
    print("2. 根据实际数据填充报告内容")
    print("3. 缺少的数据用XX代替")
