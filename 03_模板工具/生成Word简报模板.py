#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
城发投经营财务月报 - Word简报模板生成器
自动生成包含完整章节结构的Word文档模板
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_heading_style(heading, level):
    """设置标题样式"""
    # 一级标题：方正小标宋简体 22号
    if level == 1:
        heading.style = 'Heading 1'
        for run in heading.runs:
            run.font.name = '方正小标宋简体'
            run.font.size = Pt(22)
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    
    # 二级标题：黑体 16号
    elif level == 2:
        heading.style = 'Heading 2'
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(16)
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 三级标题：楷体 16号
    elif level == 3:
        heading.style = 'Heading 3'
        for run in heading.runs:
            run.font.name = '楷体'
            run.font.size = Pt(16)
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

def set_body_style(para):
    """设置正文样式"""
    para.style = 'Normal'
    for run in para.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(16)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 行间距固定值30磅
    para_format = para.paragraph_format
    para_format.line_spacing_rule = 2  # 固定值
    para_format.line_spacing = Pt(30)
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(0)

def create_word_template():
    """创建Word简报模板"""
    
    doc = Document()
    
    # 设置文档默认样式
    style = doc.styles['Normal']
    style.font.name = '仿宋_GB2312'
    style.font.size = Pt(16)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # ========== 封面 ==========
    # 标题
    title = doc.add_heading('城发投经营财务月报', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '方正小标宋简体'
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 82, 204)  # 蓝色
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    
    # 期数
    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_text = period.add_run('2026年X月')
    period_text.font.name = '微软雅黑'
    period_text.font.size = Pt(24)
    period_text.font.color.rgb = RGBColor(0, 0, 0)
    
    # 编制单位
    unit = doc.add_paragraph()
    unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unit_text = unit.add_run('编制单位：综合管理部')
    unit_text.font.name = '仿宋_GB2312'
    unit_text.font.size = Pt(16)
    unit_text._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 编制日期
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_text = date.add_run('编制日期：2026年X月X日')
    date_text.font.name = '仿宋_GB2312'
    date_text.font.size = Pt(16)
    date_text._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    toc_heading = doc.add_heading('目录', level=1)
    set_heading_style(toc_heading, 1)
    
    toc_items = [
        '一、本月经营概览',
        '二、PPP项目板块',
        '三、供热项目板块',
        '四、平台公司板块',
        '五、本部板块',
        '六、风险预警与建议'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        set_body_style(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ========== 第一章：本月经营概览 ==========
    ch1_heading = doc.add_heading('一、本月经营概览', level=1)
    set_heading_style(ch1_heading, 1)
    
    # （一）核心指标完成情况
    sec1_1 = doc.add_heading('（一）核心指标完成情况', level=2)
    set_heading_style(sec1_1, 2)
    
    # 添加表格说明
    p1 = doc.add_paragraph('下表展示了本月核心经营指标的完成情况：')
    set_body_style(p)
    
    # 创建表格（占位符）
    table1 = doc.add_table(rows=2, cols=6)
    table1.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '指标'
    hdr_cells[1].text = '本月实际'
    hdr_cells[2].text = '同比增长'
    hdr_cells[3].text = '环比增长'
    hdr_cells[4].text = '年度目标'
    hdr_cells[5].text = '完成进度'
    
    # 数据行（占位符）
    data_cells = table1.rows[1].cells
    data_cells[0].text = '示例：营业收入'
    data_cells[1].text = '18.82亿元'
    data_cells[2].text = '+5.2%'
    data_cells[3].text = '+2.1%'
    data_cells[4].text = '45.0亿元'
    data_cells[5].text = '41.8%'
    
    doc.add_paragraph()  # 空行
    
    # （二）板块利润贡献分析
    sec1_2 = doc.add_heading('（二）板块利润贡献分析', level=2)
    set_heading_style(sec1_2, 2)
    
    p2 = doc.add_paragraph('[插入饼图：PPP项目65%、供热项目20%、平台公司10%、本部5%]')
    set_body_style(p2)
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after = Pt(6)
    
    # （三）本月亮点
    sec1_3 = doc.add_heading('（三）本月亮点', level=2)
    set_heading_style(sec1_3, 2)
    
    highlights = [
        'PPP项目回款创历史新高：本月实现回款1.13亿元，创单月回款额新高。',
        '供热项目成本控制显著：单位面积成本同比下降5%，成本控制成效明显。',
        '平台公司管理提升明显：3家公司利润同比增长超过20%，管理提升效果显著。'
    ]
    
    for highlight in highlights:
        p = doc.add_paragraph(highlight, style='List Number')
        set_body_style(p)
    
    # （四）本月关注
    sec1_4 = doc.add_heading('（四）本月关注', level=2)
    set_heading_style(sec1_4, 2)
    
    concerns = [
        'PPP项目应收账款余额较高：应收账款余额58.84亿元，需加大清收力度。',
        '供热项目用户满意度下降：用户满意度评分下降0.5分，需改进服务质量。',
        '本部管理费用控制需加强：部分费用超预算，需加强费用控制。'
    ]
    
    for concern in concerns:
        p = doc.add_paragraph(concern, style='List Number')
        set_body_style(p)
    
    doc.add_page_break()
    
    # ========== 第二章：PPP项目板块 ==========
    ch2_heading = doc.add_heading('二、PPP项目板块', level=1)
    set_heading_style(ch2_heading, 1)
    
    # （一）核心指标
    sec2_1 = doc.add_heading('（一）核心指标', level=2)
    set_heading_style(sec2_1, 2)
    
    p3 = doc.add_paragraph('[插入核心指标表格 - 包含营业收入、利润总额、回款金额、应收账款余额等]')
    set_body_style(p3)
    
    # （二）项目状态分析
    sec2_2 = doc.add_heading('（二）项目状态分析', level=2)
    set_heading_style(sec2_2, 2)
    
    p4 = doc.add_paragraph('正常运作项目：15个')
    set_body_style(p4)
    p4 = doc.add_paragraph('预警项目：3个')
    set_body_style(p4)
    p4 = doc.add_paragraph('问题项目：1个')
    set_body_style(p4)
    
    p5 = doc.add_paragraph('[插入项目状态分布饼图]')
    set_body_style(p5)
    
    # （三）效益排行
    sec2_3 = doc.add_heading('（三）效益排行', level=2)
    set_heading_style(sec2_3, 2)
    
    p6 = doc.add_paragraph('利润贡献前5名项目：')
    set_body_style(p6)
    
    top5 = [
        '1. 许昌项目：0.25亿元',
        '2. 南阳项目：0.18亿元',
        '3. 安阳项目：0.15亿元',
        '4. 鹤壁项目：0.12亿元',
        '5. 濮阳项目：0.10亿元'
    ]
    
    for item in top5:
        p = doc.add_paragraph(item, style='List Number')
        set_body_style(p)
    
    doc.add_page_break()
    
    # ========== 第三章：供热项目板块 ==========
    ch3_heading = doc.add_heading('三、供热项目板块', level=1)
    set_heading_style(ch3_heading, 1)
    
    # （一）核心指标
    sec3_1 = doc.add_heading('（一）核心指标', level=2)
    set_heading_style(sec3_1, 2)
    
    p7 = doc.add_paragraph('[插入核心指标表格 - 包含供热面积、营业收入、利润总额、单位面积利润等]')
    set_body_style(p7)
    
    # （二）区域分布
    sec3_2 = doc.add_heading('（二）区域分布', level=2)
    set_heading_style(sec3_2, 2)
    
    p8 = doc.add_paragraph('[插入地图热力图：显示各区域供热面积和利润贡献]')
    set_body_style(p8)
    
    # （三）成本分析
    sec3_3 = doc.add_heading('（三）成本分析', level=2)
    set_heading_style(sec3_3, 2)
    
    p9 = doc.add_paragraph('成本构成：')
    set_body_style(p9)
    
    cost_items = [
        '燃料成本：60%',
        '人工成本：25%',
        '维护成本：10%',
        '其他成本：5%'
    ]
    
    for item in cost_items:
        p = doc.add_paragraph(item, style='List Bullet')
        set_body_style(p)
    
    p10 = doc.add_paragraph('[插入成本构成饼图]')
    set_body_style(p10)
    
    doc.add_page_break()
    
    # ========== 第四章：平台公司板块 ==========
    ch4_heading = doc.add_heading('四、平台公司板块', level=1)
    set_heading_style(ch4_heading, 1)
    
    # （一）核心指标
    sec4_1 = doc.add_heading('（一）核心指标', level=2)
    set_heading_style(sec4_1, 2)
    
    p11 = doc.add_paragraph('[插入核心指标表格 - 包含公司数量、营业收入、利润总额、净资产等]')
    set_body_style(p11)
    
    # （二）效益排行
    sec4_2 = doc.add_heading('（二）效益排行', level=2)
    set_heading_style(sec4_2, 2)
    
    p12 = doc.add_paragraph('利润贡献前3名公司：')
    set_body_style(p12)
    
    top3 = [
        '1. 新郑公司：0.10亿元',
        '2. 南阳公司：0.08亿元',
        '3. 漯河公司：0.05亿元'
    ]
    
    for item in top3:
        p = doc.add_paragraph(item, style='List Number')
        set_body_style(p)
    
    doc.add_page_break()
    
    # ========== 第五章：本部板块 ==========
    ch5_heading = doc.add_heading('五、本部板块', level=1)
    set_heading_style(ch5_heading, 1)
    
    # （一）核心指标
    sec5_1 = doc.add_heading('（一）核心指标', level=2)
    set_heading_style(sec5_1, 2)
    
    p13 = doc.add_paragraph('[插入核心指标表格 - 包含管理费用、财务费用、人员数量、人均费用等]')
    set_body_style(p13)
    
    # （二）费用控制分析
    sec5_2 = doc.add_heading('（二）费用控制分析', level=2)
    set_heading_style(sec5_2, 2)
    
    p14 = doc.add_paragraph('管理费用构成：')
    set_body_style(p14)
    
    expense_items = [
        '职工薪酬：60%',
        '办公费：15%',
        '业务招待费：10%',
        '其他费用：15%'
    ]
    
    for item in expense_items:
        p = doc.add_paragraph(item, style='List Bullet')
        set_body_style(p)
    
    p15 = doc.add_paragraph('[插入管理费用构成饼图]')
    set_body_style(p15)
    
    doc.add_page_break()
    
    # ========== 第六章：风险预警与建议 ==========
    ch6_heading = doc.add_heading('六、风险预警与建议', level=1)
    set_heading_style(ch6_heading, 1)
    
    # （一）风险预警
    sec6_1 = doc.add_heading('（一）风险预警', level=2)
    set_heading_style(sec6_1, 2)
    
    # 风险表格（占位符）
    table2 = doc.add_table(rows=2, cols=6)
    table2.style = 'Light Grid Accent 1'
    
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '风险点'
    hdr_cells2[1].text = '风险等级'
    hdr_cells2[2].text = '风险描述'
    hdr_cells2[3].text = '建议措施'
    hdr_cells2[4].text = '责任部门'
    hdr_cells2[5].text = '完成时限'
    
    data_cells2 = table2.rows[1].cells
    data_cells2[0].text = '示例：PPP项目应收账款余额较高'
    data_cells2[1].text = '🔴高'
    data_cells2[2].text = '应收账款余额58.84亿元，回收周期长'
    data_cells2[3].text = '加大清收力度，建立双周例会制度'
    data_cells2[4].text = '资产管理部'
    data_cells2[5].text = '持续推进'
    
    doc.add_paragraph()  # 空行
    
    # （二）下月重点工作建议
    sec6_2 = doc.add_heading('（二）下月重点工作建议', level=2)
    set_heading_style(sec6_2, 2)
    
    suggestions = [
        '加大PPP项目回款力度：抓住回款黄金期，确保完成全年回款目标。',
        '改进供热项目服务质量：针对用户反馈问题，制定改进措施，提升用户满意度。',
        '加强本部费用控制：严格执行费用预算，控制不必要支出。',
        '推进轻资产项目落地：加快轻资产项目储备和落地，提升资产收益率。'
    ]
    
    for suggestion in suggestions:
        p = doc.add_paragraph(suggestion, style='List Number')
        set_body_style(p)
    
    doc.add_page_break()
    
    # ========== 签名栏 ==========
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    sig_text1 = sig_para.add_run('编制人：')
    sig_text1.font.name = '仿宋_GB2312'
    sig_text1.font.size = Pt(16)
    sig_text1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    sig_para2 = doc.add_paragraph()
    sig_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    sig_text2 = sig_para2.add_run('审核人：')
    sig_text2.font.name = '仿宋_GB2312'
    sig_text2.font.size = Pt(16)
    sig_text2._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    sig_para3 = doc.add_paragraph()
    sig_para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    sig_text3 = sig_para3.add_run('批准人：')
    sig_text3.font.name = '仿宋_GB2312'
    sig_text3.font.size = Pt(16)
    sig_text3._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 保存文件
    output_file = '城发投经营财务月报模板.docx'
    doc.save(output_file)
    
    print(f"✅ Word模板已生成：{output_file}")
    print(f"📄 共生成 {len(doc.paragraphs)} 个段落")
    print("\n📋 文档结构说明：")
    print("   封面：标题、期数、编制单位、编制日期")
    print("   目录：自动生成目录（需在Word中更新域）")
    print("   第一章：本月经营概览（4节）")
    print("   第二章：PPP项目板块（3节）")
    print("   第三章：供热项目板块（3节）")
    print("   第四章：平台公司板块（2节）")
    print("   第五章：本部板块（2节）")
    print("   第六章：风险预警与建议（2节）")
    print("   签名栏：编制人、审核人、批准人")
    print("\n💡 使用说明：")
    print("   1. 打开生成的.docx文件")
    print("   2. 按Ctrl+A全选，按F9更新所有域（目录、页码等）")
    print("   3. 替换占位符内容为实际数据和图表")
    print("   4. 根据实际需求调整内容和格式")
    print("   5. 另存为.dotx文件作为模板使用")
    print("\n🎨 格式说明：")
    print("   一级标题：方正小标宋简体 22号 不加粗")
    print("   二级标题：黑体 16号 不加粗")
    print("   三级标题：楷体 16号 不加粗")
    print("   正文：仿宋_GB2312 16号 不加粗")
    print("   行间距：固定值30磅")
    print("   段前段后：0行")
    print("   字体颜色：全部黑色")

if __name__ == "__main__":
    create_word_template()
