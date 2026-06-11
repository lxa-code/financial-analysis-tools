# -*- coding: utf-8 -*-
"""生成郑州区域材料质量审核报告 Word 文件（公文排版）"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# 页面设置：A4，标准公文边距
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
style.font.name = '仿宋'
style.font.size = Pt(16)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

# 标题样式
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.line_spacing = 1.5
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(12)
    elif i == 2:
        hs.font.size = Pt(18)
        hs.paragraph_format.line_spacing = 1.5
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(6)
    else:
        hs.font.size = Pt(16)
        hs.paragraph_format.line_spacing = 1.5
        hs.paragraph_format.space_before = Pt(3)
        hs.paragraph_format.space_after = Pt(3)


def add_para(text, bold=False, font_name='仿宋', font_size=Pt(16), alignment=None,
             first_line_indent=Cm(0.74), space_before=Pt(0), space_after=Pt(0), color=None):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return p


def add_heading_text(text, level=1):
    """添加标题段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    if level == 0:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        font_size = Pt(24)
        font_name = '方正小标宋简体'
        bold = False
    elif level == 1:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0.74)
        font_size = Pt(16)
        font_name = '黑体'
        bold = True
    elif level == 2:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0.74)
        font_size = Pt(16)
        font_name = '楷体'
        bold = True
    else:
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0.74)
        font_size = Pt(16)
        font_name = '仿宋'
        bold = False
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def set_cell_font(cell, text, font_name='仿宋', font_size=Pt(12), bold=False, alignment=None):
    """设置表格单元格文字"""
    # 清除默认段落
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def shade_cells(row, color="D5E8F0"):
    """给整行设置底色"""
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def add_table_with_data(headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_font(hdr.cells[i], h, font_name='黑体', font_size=Pt(12), bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cells(hdr, "1F4E79")
    # 表头文字白色
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            set_cell_font(row.cells[c_idx], str(val), font_size=Pt(11),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT)
        if r_idx % 2 == 1:
            shade_cells(row, "EAF0F6")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table


# ============================================================
# 正文内容
# ============================================================

# 主标题
add_heading_text('郑州区域材料质量审核报告', level=0)

# 一、材料清单概览
add_heading_text('一、材料清单概览', level=1)

add_para('本次审核共涉及郑州区域六家项目公司提交的六份材料，文件清单如下：')

headers1 = ['序号', '文件名', '文档类型', '核心内容']
rows1 = [
    ['1', '新郑城发.docx', '重点工作计划+预算说明', '特许经营权获取、市场拓展、供热检修、章程修订、账款回收'],
    ['2', '汇融热力.docx', '重点工作计划+预算说明', '轻资产运营、供热"一张网"、ABS金融工具、品牌生态'],
    ['3', '河南启迪零碳.docx', '预算汇报', '地源热泵供能、补贴依赖、诉讼欠款'],
    ['4', '河南城发生态.docx', '重点工作计划', '会展路BT项目结算审计、回款化债、移交难题'],
    ['5', '航空港热力.docx', '全面预算+经营分析', '供热运营+安装业务、供热企业整合、管网建设'],
    ['6', '郑州环保能源河南分公司.docx', '重点工作计划+预算', '轻资产运营项目拓展、工程建设、风险分析'],
]
add_table_with_data(headers1, rows1, [1.0, 4.0, 3.5, 6.5])

# 二、整体评价
add_heading_text('二、整体评价', level=1)

add_para('总体判断：材料基本齐备，但质量参差不齐，缺乏统一规范。', bold=True)
add_para('六份材料中，航空港热力最为完整规范，包含预算编制报告和经营分析报告两部分完整文档；新郑城发、汇融热力、河南分公司结构较为清晰但细节不足；河南启迪零碳、河南城发生态最为薄弱。主要问题集中在：格式不统一、时间节点普遍过期、问题清单缺失、"需要领导决策事项"过于笼统、预算数据审慎性不足等五个方面。')

# 三、逐份材料审核意见
add_heading_text('三、逐份材料审核意见', level=1)

# --- (一) 新郑城发 ---
add_heading_text('（一）新郑城发.docx', level=2)
add_para('1. 优点', bold=True)
add_para('重点工作计划结构清晰，"五大任务"逻辑完整；预算说明部分数据较为详细，有同期对比分析；风险分析按"战略—经营—财务—运营—专项"五维度分类，框架合理。')

add_para('2. 不足与改进建议', bold=True)
headers_n = ['序号', '存在问题', '改进建议']
rows_xz = [
    ['1', '缺少问题清单：标题虽有"经营发展问题清单"，但正文未列出明确的问题条目',
     '在"公司现存问题分析"后增设"问题清单表"，按编号—问题—严重程度—影响—对策五列呈现'],
    ['2', '特许经营权时间节点过于乐观：计划5月30日取得主导权、7月31日完成评估，但当前已是6月',
     '更新实际进展，标注已滞后节点及原因，重新排期'],
    ['3', '"需要领导决策事项"过于笼统：仅写"协助对接相关政府部门"',
     '明确需要哪个层级领导、对接哪个具体部门、什么时间节点前需要完成'],
    ['4', '预算报告和重点工作计划合并在一个文件，目录显示有12个章节但正文仅部分内容',
     '建议分拆为两个独立文件，或确保完整覆盖目录所列全部章节'],
    ['5', '回款催收缺少量化目标分解：仅说"至少支付800万元"，缺少月度催收计划',
     '补充按季度/月度的回款分解计划，明确各阶段催收责任人'],
]
add_table_with_data(headers_n, rows_xz, [0.8, 6.5, 7.7])

# --- (二) 汇融热力 ---
add_heading_text('（二）汇融热力.docx', level=2)
add_para('1. 优点', bold=True)
add_para('战略定位清晰，"科技+金融"双轮驱动的定位表述有力；营收与新签合同目标分解到季度，有时间节奏；八大重点工作覆盖全面，从轻资产运营到人才建设均有涉及。')

add_para('2. 不足与改进建议', bold=True)
rows_hr = [
    ['1', '营收1亿元目标缺乏充分支撑：收入构成表中综合能源服务4800万、集中采购2200万等，但具体项目落地路径未说明',
     '每个板块需补充"项目—金额—当前进展—关键里程碑—风险"的明细表'],
    ['2', '新签合同5亿元过于宏大：包含巨鹿、宁陵、开封等多个未落地项目，实际签约概率未评估',
     '按"已签约/谈判中/意向"三级分类，给出每个项目的签约概率及最晚签约时间'],
    ['3', '人才均为兼职人员：全部为兼职人员，无人员在本公司取薪——这对一个要做1亿营收的平台公司来说是重大治理缺陷',
     '说明核心岗位专职化时间表，至少明确总经理、财务负责人、技术负责人的到岗计划'],
    ['4', '预算编制说明和重点工作计划分属两个文体，中间缺少过渡',
     '统一格式，或在目录中明确区分'],
    ['5', 'ABS金融工具方案描述过于简略',
     '补充操作流程图、各参与方职责、时间表、已对接的金融机构清单'],
]
add_table_with_data(headers_n, rows_hr, [0.8, 6.5, 7.7])

# --- (三) 河南启迪零碳 ---
add_heading_text('（三）河南启迪零碳.docx', level=2)
add_para('1. 优点', bold=True)
add_para('定位清晰，专注于地源热泵综合能源服务；风险分析较为坦诚，承认了账户冻结、欠款等现实问题。')

add_para('2. 不足与改进建议', bold=True)
rows_qd = [
    ['1', '仅有预算汇报，缺少重点工作计划：与其他5份材料格式不一致',
     '补充2026年重点工作计划，至少包含：诉讼推进、账户解冻、市场拓展、设备运维四个板块'],
    ['2', '供能补贴收入断崖式下降：营收从475.63万降至332.3万（降30%），利润从197.6万降至72.59万（降63%），但应对措施仅一句"拓展本园区和周边的供能面积"',
     '补充详细的增收方案：具体哪些周边建筑、预计供能面积、何时签约、预计增收金额'],
    ['3', '启迪东龙累计欠款1084.54万元，这是致命性问题但着墨太少',
     '补充诉讼进展时间表、财产线索排查情况、胜诉后执行可行性评估、如无法执行对公司持续经营的影响'],
    ['4', '风险分类编号跳跃：出现"（五）各环节专项风险"，缺少"（四）"',
     '修正编号，确保连续性'],
]
add_table_with_data(headers_n, rows_qd, [0.8, 6.5, 7.7])

# --- (四) 河南城发生态 ---
add_heading_text('（四）河南城发生态.docx', level=2)
add_para('1. 优点', bold=True)
add_para('聚焦BT项目结算审计、回款、移交三大核心，方向明确；回款部分对困难分析较为客观（政府隐形债、财政紧张等）。')

add_para('2. 不足与改进建议', bold=True)
rows_cf = [
    ['1', '文件最短，仅约2页，是所有材料中最薄弱的',
     '大幅扩充，至少补充：项目背景介绍、历年回款明细表、当前应收款账龄结构、各化债路径的可行性对比'],
    ['2', '缺少预算/财务数据：没有2026年预算目标、没有资金平衡计划',
     '补充2026年预计收入、成本、利润、现金流等基本财务指标'],
    ['3', '缺少风险分析：没有对审计不通过、回款持续无进展等情景的压力测试',
     '补充风险矩阵，至少包含"乐观/基准/悲观"三种情景下的现金流预测'],
    ['4', '"需要领导决策事项"缺失：文中提到"必要时请公司领导进行高层协调"但未明确具体诉求',
     '增设"需要集团支持事项"板块，明确对接层级、对接部门、时间节点'],
    ['5', '时间节点模糊："计划在4月中旬""计划在4月底前"，当前已6月',
     '更新所有时间节点为实际进展状态，标注已完成/滞后/取消'],
]
add_table_with_data(headers_n, rows_cf, [0.8, 6.5, 7.7])

# --- (五) 航空港热力 ---
add_heading_text('（五）航空港热力.docx', level=2)
add_para('1. 优点', bold=True)
add_para('六份材料中质量最高。包含预算编制报告和经营分析报告两份完整文档；财务数据翔实，有5年趋势对比（2021—2025），资产、负债、收入、利润、现金流均有详细分析；供热面积明细表非常详细，33个安置区+23个商业小区逐一列出；生产指标（热耗、电耗、水耗、室温合格率等）逐年对比，体现精细化运营。')

add_para('2. 不足与改进建议', bold=True)
rows_hk = [
    ['1', '两份报告存在数据不一致：经营分析报告中2026年净利润39.14万元，预算编制报告中净利润55.61万元',
     '核实并统一口径，标注差异原因'],
    ['2', '缺少独立的问题清单：问题散落在经营分析报告第六部分，但未以结构化清单呈现',
     '提取为"2026年重点问题清单"，按问题—现状—影响—对策—责任人—完成时限六列格式化'],
    ['3', '供热业务持续亏损：2025年供热毛利仅103.21万元（扣除折旧后亏损-2232万），2026年预算供热毛利2330万但扣除折旧后仍亏损-182万',
     '补充供热业务盈利拐点分析：预计何时实现供热业务自身盈亏平衡'],
    ['4', '安装业务毛利骤降：从2025年的5495.83万降至2026年预算的3870.50万（降29.57%），原因分析不够充分',
     '细化城中村改造项目收入确认政策的变更影响，补充在手订单及新签合同预期'],
]
add_table_with_data(headers_n, rows_hk, [0.8, 6.5, 7.7])

# --- (六) 河南分公司 ---
add_heading_text('（六）郑州环保能源河南分公司.docx', level=2)
add_para('1. 优点', bold=True)
add_para('四个项目拓展方向清晰，每个都有经济分析；风险分析非常详尽，按战略—经营—财务—运营—专项五维度展开；优化建议"创利—创现—创值"三维度框架与集团战略一致。')

add_para('2. 不足与改进建议', bold=True)
rows_hn = [
    ['1', '与汇融热力内容高度重叠：预算数据几乎完全一致（都是营收2516.79万、利润157.61万），但属于两个不同的法律主体',
     '明确两个主体的关系（是否为"一套人马两块牌子"），说明业务划分边界，避免重复核算'],
    ['2', '项目拓展多为"谋划事宜"：巨鹿、襄城蒸汽、暖气片安装、高校合同能源管理等均处于早期意向阶段',
     '按"已签约/谈判中/意向接触"分级，标注各项目的成熟度和预计签约时间'],
    ['3', '人员模式存疑：兼职人员共5名但分布在三个不同法人主体，权责不清',
     '说明兼职人员的管理关系、考核机制、薪酬承担方式'],
    ['4', '章节编号混乱：优化建议标为"十四"但前面只有十个章节',
     '统一编号体系'],
]
add_table_with_data(headers_n, rows_hn, [0.8, 6.5, 7.7])

# 四、共性问题
add_heading_text('四、共性问题（六份材料的通病）', level=1)

add_para('（一）格式不统一，缺乏规范化模板。', bold=True)
add_para('六份材料格式各异：有的含目录、有的没有；有的有预算+计划两份内容，有的只有一份；有的用"一、（一）、1"编号，有的混用。建议由要素保障组制定统一的《重点工作计划及问题清单编制模板》，明确以下必含模块：公司概况、2026年度财务预算目标（一张表）、问题清单（编号—问题—严重程度—影响—对策—责任人—时限）、重点工作计划（每项含目标—举措—里程碑—责任人）、需要集团协调/决策事项、风险分析（战略—经营—财务—运营—专项）。')

add_para('（二）时间节点普遍过期或缺失。', bold=True)
add_para('多份材料标注的"4月""5月"时间节点已经过去（当前6月4日），但材料未更新实际进展。建议所有材料统一增加"当前进展"列，标注已完成/进行中/滞后的实际状态。')

add_para('（三）"需要领导决策事项"过于笼统。', bold=True)
add_para('多数写成"协助对接政府部门""请领导协调"等空泛表述。建议统一为"事项—对接部门—对接层级—期望结果—最晚时间"五要素。')

add_para('（四）问题清单缺失或不规范。', bold=True)
add_para('仅河南分公司有较系统的风险分析，其他材料的问题散落在各处。建议每份材料必须前置"2026年重点问题清单"，作为后续工作计划的前提和依据。')

add_para('（五）缺乏跨公司协同视角。', bold=True)
add_para('新郑城发与航空港热力有管输业务往来、汇融热力与河南分公司业务重叠，但材料中缺乏"需要兄弟公司协同的事项"板块。建议增设"跨公司协同事项"板块，明确上下游衔接关系。')

add_para('（六）预算数据的审慎性需加强。', bold=True)
add_para('部分预算目标设定偏乐观（如汇融热力营收1亿、新签合同5亿，但团队全为兼职），部分数据在不同文件间不一致（如航空港热力两份报告净利润差16.47万）。建议要素保障组对关键预算数据进行交叉核验。')

# 五、综合评分
add_heading_text('五、综合评分', level=1)

add_para('按照结构完整性、数据翔实度、问题清晰度、可操作性、格式规范性五个维度，对六份材料进行综合评分（A—D级），结果如下：')

headers_s = ['维度', '新郑城发', '汇融热力', '启迪零碳', '城发生态', '航空港热力', '河南分公司', '平均']
rows_s = [
    ['结构完整性', 'B', 'B+', 'C', 'D', 'A', 'B', 'B-'],
    ['数据翔实度', 'B+', 'B', 'B-', 'D', 'A', 'B+', 'B'],
    ['问题清晰度', 'C+', 'C+', 'B-', 'C', 'B', 'A-', 'B-'],
    ['可操作性', 'B-', 'C+', 'C-', 'C-', 'B+', 'B-', 'C+'],
    ['格式规范性', 'B', 'B', 'C+', 'C-', 'A-', 'B-', 'B-'],
    ['综合', 'B-', 'B-', 'C+', 'D+', 'A-', 'B-', 'C+'],
]
add_table_with_data(headers_s, rows_s, [2.5, 1.6, 1.6, 1.6, 1.6, 1.6, 1.6, 1.6])

# 六、优先改进建议
add_heading_text('六、优先改进建议（Top 5）', level=1)

add_para('第一，制定统一模板，要求所有项目公司按统一格式重新提报。模板须含问题清单、工作计划、预算目标、需协调事项、风险分析五大板块。')
add_para('第二，更新所有时间节点至当前实际状态，标注已完成、滞后或取消的原因及重新排期。')
add_para('第三，补齐河南城发生态的财务数据，当前该文件完全没有预算内容，应至少补充2026年预计收入、成本、利润、现金流等基本指标。')
add_para('第四，核实并统一航空港热力两份报告间的数据矛盾（净利润差16.47万），查明差异原因并修正。')
add_para('第五，要求汇融热力和河南分公司明确两个主体的业务边界、人员归属及核算关系，避免重复核算和权责不清。')

# 落款
doc.add_paragraph()
add_para('', first_line_indent=None)
add_para('要素保障组', bold=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=None)
add_para('2026年6月4日', bold=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=None)

# ============================================================
# 保存文件
# ============================================================
output_path = r'd:\BaiduSyncdisk\CODE\城发投资工作专区\杂项\重点工作计划及问题清单\郑州区域\郑州区域材料质量审核报告.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'报告已生成: {output_path}')
