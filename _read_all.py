import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx, openpyxl

# 日报 - 只看日期标题
docx_path = r"e:\BaiduSyncdisk\CODE\城发投资工作专区\10_日常工作\个人工作日报\新建 DOCX 文档.docx"
doc = docx.Document(docx_path)
dates = []
for para in doc.paragraphs:
    t = para.text.strip()
    if t and ('202' in t or '月' in t) and len(t) < 30:
        dates.append(t)
print("日报中的日期:")
for d in dates:
    print(f"  {d}")
print(f"\n共 {len(dates)} 个日期标记")

# 输出全部日报内容
print("\n" + "=" * 60)
print("【日报全部内容】")
print("=" * 60)
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t:
        print(f"[{i}] {t}")
